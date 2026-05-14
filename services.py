from docx import Document
import os, re


def copiar_estilo(origem, destino):

    destino.bold = origem.bold
    destino.italic = origem.italic
    destino.underline = origem.underline

    destino.font.name = origem.font.name
    destino.font.size = origem.font.size
    destino.font.bold = origem.font.bold
    destino.font.italic = origem.font.italic

    try:
        destino.font.color.rgb = origem.font.color.rgb
    except:
        pass

def substituir_runs(paragraph, substituicoes):

    for run in paragraph.runs:

        for buscar, substituir in substituicoes.items():

            if buscar in run.text:

                run.text = run.text.replace(
                    buscar,
                    str(substituir)
                )

def aplicar_substituicoes(doc, substituicoes):

    # Parágrafos
    for p in doc.paragraphs:
        substituir_runs(p, substituicoes)

    # Tabelas
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for p in celula.paragraphs:
                    substituir_runs(p, substituicoes)

    # Header / Footer
    for secao in doc.sections:

        for p in secao.header.paragraphs:
            substituir_runs(p, substituicoes)

        for p in secao.footer.paragraphs:
            substituir_runs(p, substituicoes)

    return doc


def substituir_textos_documento(
    caminho_modelo,
    substituicoes
):

    if not os.path.exists(caminho_modelo):
        raise Exception("Arquivo não encontrado")

    doc = Document(caminho_modelo)

    doc = aplicar_substituicoes(
        doc,
        substituicoes
    )

    return doc


def preencher_documento(
    caminho_modelo,
    dados
):

    substituicoes = {
        f"{{{{{k}}}}}": v
        for k, v in dados.items()
    }

    return substituir_textos_documento(
        caminho_modelo,
        substituicoes
    )

def extrair_placeholders(caminho_modelo):
    if not os.path.exists(caminho_modelo):
        raise Exception("Arquivo não encontrado")
    
    doc = Document(caminho_modelo)
    textos = []

    #Parágrafos
    for p in doc.paragraphs:
        textos.append(p.text)

    #Tabelas
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for p in celula.paragraphs:
                    textos.append(p.text)

    #Header/Footer
    for secao in doc.sections:
        for p in secao.header.paragraphs:
            textos.append(p.text)

        for p in secao.footer.paragraphs:
            textos.append(p.text)

    texto_completo = "\n".join(textos)
    encontrados = re.findall(
        r"\{\{(.*?)\}\}",
        texto_completo
    )
    #remove duplicados
    placeholders = list(set(encontrados))
    return placeholders